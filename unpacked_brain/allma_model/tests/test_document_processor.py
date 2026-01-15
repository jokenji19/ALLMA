import unittest
import os
import tempfile
from docx import Document
from openpyxl import Workbook
from PIL import Image
import numpy as np
from allma_model.core.document_processor import DocumentProcessor

class TestDocumentProcessor(unittest.TestCase):
    def setUp(self):
        """Setup per i test"""
        self.processor = DocumentProcessor()
        self.test_dir = tempfile.mkdtemp()
        
    def tearDown(self):
        """Pulizia dopo i test"""
        for root, dirs, files in os.walk(self.test_dir, topdown=False):
            for name in files:
                os.remove(os.path.join(root, name))
            for name in dirs:
                os.rmdir(os.path.join(root, name))
        os.rmdir(self.test_dir)
        
    def test_process_text_file(self):
        """Test processamento file di testo"""
        # Crea un file di testo
        content = "Hello ALLMA!\nThis is a test."
        file_path = os.path.join(self.test_dir, "test.txt")
        with open(file_path, 'w') as f:
            f.write(content)
            
        # Processa il file
        result = self.processor.process_document(file_path)
        
        # Verifica i risultati
        self.assertIn('content', result)
        self.assertEqual(result['content'], content)
        self.assertEqual(len(result['lines']), 2)
        self.assertEqual(result['lines'][0], "Hello ALLMA!")
        
    def test_process_docx(self):
        """Test processamento file Word"""
        # Crea un documento Word
        doc = Document()
        doc.add_paragraph("Hello ALLMA!")
        doc.add_paragraph("This is a test.")
        
        # Aggiungi una tabella
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        
        # Salva il documento
        file_path = os.path.join(self.test_dir, "test.docx")
        doc.save(file_path)
        
        # Processa il file
        result = self.processor.process_document(file_path)
        
        # Verifica i risultati
        self.assertIn('paragraphs', result)
        self.assertEqual(len(result['paragraphs']), 2)
        self.assertEqual(result['paragraphs'][0]['text'], "Hello ALLMA!")
        self.assertIn('tables', result)
        self.assertEqual(len(result['tables']), 1)
        self.assertEqual(result['tables'][0][0][0], "A1")
        
    def test_process_excel(self):
        """Test processamento file Excel"""
        # Crea un workbook Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "Test Sheet"
        
        # Aggiungi dati
        ws['A1'] = "Hello"
        ws['B1'] = "ALLMA"
        ws['A2'] = 123
        ws['B2'] = "Test"
        
        # Salva il file
        file_path = os.path.join(self.test_dir, "test.xlsx")
        wb.save(file_path)
        
        # Processa il file
        result = self.processor.process_document(file_path)
        
        # Verifica i risultati
        self.assertIn('sheets', result)
        self.assertEqual(len(result['sheets']), 1)
        self.assertEqual(result['sheets'][0]['name'], "Test Sheet")
        self.assertEqual(result['sheets'][0]['data'][0][0], "Hello")
        self.assertEqual(result['sheets'][0]['data'][1][1], "Test")
        
    def test_process_image(self):
        """Test processamento immagine"""
        # Crea un'immagine con testo
        width = 400
        height = 100
        image = np.full((height, width, 3), 255, dtype=np.uint8)
        
        # Salva l'immagine
        file_path = os.path.join(self.test_dir, "test.png")
        Image.fromarray(image).save(file_path)
        
        # Processa l'immagine
        result = self.processor.process_document(file_path)
        
        # Verifica i risultati
        self.assertIn('text', result)
        self.assertIn('statistics', result)
        
    def test_extract_text(self):
        """Test estrazione testo da vari formati"""
        # Crea un file di testo
        text_content = "Hello ALLMA!\nThis is a test."
        text_path = os.path.join(self.test_dir, "test.txt")
        with open(text_path, 'w') as f:
            f.write(text_content)
            
        # Estrai il testo
        text = self.processor.extract_text_from_document(text_path)
        
        # Verifica il risultato
        self.assertIn("Hello ALLMA!", text)
        self.assertIn("This is a test", text)
        
    def test_unsupported_format(self):
        """Test gestione formato non supportato"""
        # Crea un file con estensione non supportata
        file_path = os.path.join(self.test_dir, "test.xyz")
        with open(file_path, 'w') as f:
            f.write("test")
            
        # Prova a processare il file
        result = self.processor.process_document(file_path)
        
        # Verifica che ci sia un errore
        self.assertIn('error', result)
        
    def test_file_not_found(self):
        """Test gestione file non esistente"""
        file_path = os.path.join(self.test_dir, "nonexistent.txt")
        result = self.processor.process_document(file_path)
        self.assertIn('error', result)
        self.assertIn('File non trovato', result['error'])
        
    def test_unsupported_format(self):
        """Test gestione formato non supportato"""
        # Crea un file con estensione non supportata
        file_path = os.path.join(self.test_dir, "test.xyz")
        with open(file_path, 'w') as f:
            f.write("test content")
            
        result = self.processor.process_document(file_path)
        self.assertIn('error', result)
        self.assertIn('Formato file non supportato', result['error'])
        
    def test_empty_text_file(self):
        """Test processamento file di testo vuoto"""
        file_path = os.path.join(self.test_dir, "empty.txt")
        with open(file_path, 'w') as f:
            f.write("")
            
        result = self.processor.process_document(file_path)
        self.assertIn('content', result)
        self.assertEqual(result['content'], "")
        self.assertEqual(len(result['lines']), 0)
        
    def test_image_metadata(self):
        """Test estrazione metadati immagine"""
        # Crea un'immagine con metadati
        width, height = 400, 200
        img = Image.new('RGB', (width, height), color='white')
        file_path = os.path.join(self.test_dir, "test_metadata.jpg")
        img.save(file_path, quality=95, optimize=True)
        
        result = self.processor.process_document(file_path)
        self.assertIn('metadata', result)
        self.assertIn('dimensions', result['metadata'])
        self.assertEqual(result['metadata']['dimensions'], (width, height))
        
    def test_docx_with_headers(self):
        """Test processamento Word con intestazioni"""
        doc = Document()
        
        # Aggiungi intestazione
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "Test Header"
        
        # Aggiungi contenuto normale
        doc.add_paragraph("Test Content")
        
        # Aggiungi piè di pagina
        footer = section.footer
        footer.paragraphs[0].text = "Test Footer"
        
        file_path = os.path.join(self.test_dir, "test_headers.docx")
        doc.save(file_path)
        
        result = self.processor.process_document(file_path)
        self.assertIn('headers', result)
        self.assertIn('footers', result)
        self.assertTrue(any("Test Header" in h['text'] for h in result['headers']))
        self.assertTrue(any("Test Footer" in f['text'] for f in result['footers']))
        
    def test_excel_formulas(self):
        """Test processamento Excel con formule"""
        wb = Workbook()
        ws = wb.active
        
        # Aggiungi dati e formule
        ws['A1'] = 10
        ws['A2'] = 20
        ws['A3'] = '=SUM(A1:A2)'
        ws['B1'] = '=A1*2'
        
        file_path = os.path.join(self.test_dir, "test_formulas.xlsx")
        wb.save(file_path)
        
        # Crea un altro workbook per i risultati
        wb_results = Workbook()
        ws_results = wb_results.active
        ws_results['A1'] = 10
        ws_results['A2'] = 20
        ws_results['A3'] = 30  # Risultato di SUM(A1:A2)
        ws_results['B1'] = 20  # Risultato di A1*2
        
        results_path = os.path.join(self.test_dir, "test_formulas_results.xlsx")
        wb_results.save(results_path)
        
        # Processa il file con le formule
        result = self.processor.process_document(file_path)
        self.assertIn('formulas', result)
        self.assertGreater(len(result['formulas']), 0)
        
        # Processa il file con i risultati
        result_data = self.processor.process_document(results_path)
        
        # Verifica i risultati delle formule
        formulas = {f['cell']: f for f in result['formulas']}
        self.assertEqual(formulas['A3']['formula'], '=SUM(A1:A2)')
        self.assertEqual(formulas['B1']['formula'], '=A1*2')
        
        # Verifica i dati del foglio con i risultati
        sheet_data = result_data['sheets'][0]['data']
        self.assertEqual(sheet_data[0][0], 10)  # A1
        self.assertEqual(sheet_data[1][0], 20)  # A2
        self.assertEqual(sheet_data[2][0], 30)  # A3 (risultato)
        self.assertEqual(sheet_data[0][1], 20)  # B1 (risultato)

    def test_corrupted_file(self):
        """Test gestione file corrotto"""
        # Crea un file docx corrotto
        file_path = os.path.join(self.test_dir, "corrupted.docx")
        with open(file_path, 'wb') as f:
            f.write(b'This is not a valid DOCX file')
            
        result = self.processor.process_document(file_path)
        self.assertIn('error', result)
        
    def test_large_file_handling(self):
        """Test gestione file di grandi dimensioni"""
        # Crea un file di testo grande
        file_path = os.path.join(self.test_dir, "large.txt")
        with open(file_path, 'w') as f:
            f.write("Test line\n" * 100000)  # Circa 1MB di testo
            
        result = self.processor.process_document(file_path)
        self.assertIn('content', result)
        self.assertIn('file_info', result)
        self.assertGreater(result['file_info']['size'], 900000)  # Almeno 900KB
        
    def test_special_characters(self):
        """Test gestione caratteri speciali"""
        content = "Hello ALLMA! 你好 Привет ⭐️ 🌟 ❤️\nSpecial: àèìòù"
        file_path = os.path.join(self.test_dir, "special.txt")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
        result = self.processor.process_document(file_path)
        self.assertEqual(result['content'], content)
        self.assertTrue(all(char in result['content'] for char in ['你', 'Привет', '⭐️', 'à']))

if __name__ == '__main__':
    unittest.main()

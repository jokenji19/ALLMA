import os
from typing import Dict, List, Optional, Any
from datetime import datetime
from pypdf import PdfReader
import pdfplumber
from docx import Document
from openpyxl import load_workbook
import cv2
import numpy as np
from PIL import Image
from .ocr_processor import OCRProcessor

class DocumentProcessor:
    def __init__(self):
        """Inizializza il processore di documenti"""
        self.ocr = OCRProcessor()
        self.supported_extensions = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'doc': self._process_docx,
            'xlsx': self._process_excel,
            'xls': self._process_excel,
            'txt': self._process_text,
            'jpg': self._process_image,
            'jpeg': self._process_image,
            'png': self._process_image,
        }
        
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Processa un documento e ne estrae il contenuto"""
        try:
            # Verifica che il file esista
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File non trovato: {file_path}")
                
            # Ottieni l'estensione del file
            extension = os.path.splitext(file_path)[1].lower().replace('.', '')
            
            # Verifica che l'estensione sia supportata
            if extension not in self.supported_extensions:
                raise ValueError(f"Formato file non supportato: {extension}")
                
            # Processa il documento usando la funzione appropriata
            processor = self.supported_extensions[extension]
            result = processor(file_path)
            
            # Aggiungi metadati comuni
            result.update({
                'file_info': {
                    'path': file_path,
                    'size': os.path.getsize(file_path),
                    'modified': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
                    'type': extension
                }
            })
            
            return result
            
        except Exception as e:
            print(f"Error in process_document: {e}")
            return {
                'error': str(e),
                'file_path': file_path
            }
            
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Processa un file PDF"""
        result = {
            'pages': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            # Estrai il testo e i metadati usando pypdf
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                
                # Estrai metadati
                result['metadata'] = reader.metadata if reader.metadata else {}
                
                # Usa pdfplumber per un'estrazione più accurata
                with pdfplumber.open(file_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        # Estrai testo
                        text = page.extract_text()
                        
                        # Estrai tabelle
                        tables = page.extract_tables()
                        
                        # Estrai immagini
                        images = page.images
                        
                        # Analizza il testo con OCR se necessario
                        if not text.strip():
                            # Converti la pagina in immagine
                            img = page.to_image()
                            img_path = f"{file_path}_page_{i}.png"
                            img.save(img_path)
                            
                            # Esegui OCR
                            ocr_results = self.ocr.extract_text_from_image(
                                cv2.imread(img_path)
                            )
                            
                            # Pulisci
                            os.remove(img_path)
                            
                            text = " ".join([r['text'] for r in ocr_results.get('text', [])])
                            
                        result['pages'].append({
                            'page_number': i + 1,
                            'text': text,
                            'tables': tables,
                            'images': [img for img in images]
                        })
                        
            return result
            
        except Exception as e:
            print(f"Error in _process_pdf: {e}")
            return {'error': str(e)}
            
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Processa un file Word"""
        result = {
            'paragraphs': [],
            'tables': [],
            'images': [],
            'headers': [],
            'footers': []
        }
        
        try:
            doc = Document(file_path)
            
            # Processa intestazioni e piè di pagina
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            result['headers'].append({
                                'text': paragraph.text,
                                'style': paragraph.style.name
                            })
                            
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            result['footers'].append({
                                'text': paragraph.text,
                                'style': paragraph.style.name
                            })
            
            # Processa paragrafi
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    result['paragraphs'].append({
                        'text': paragraph.text,
                        'style': paragraph.style.name
                    })
            
            # Processa tabelle
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                result['tables'].append(table_data)
                
            return result
            
        except Exception as e:
            print(f"Error in _process_docx: {e}")
            return {'error': str(e)}
            
    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Processa un file Excel"""
        result = {
            'sheets': [],
            'formulas': []
        }
        
        try:
            # Carica il workbook due volte: una per le formule e una per i risultati
            wb = load_workbook(file_path, data_only=False)
            wb_data = load_workbook(file_path, data_only=True)
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_data = wb_data[sheet_name]
                
                data = []
                for row in sheet_data.iter_rows():
                    # Converti i valori in stringhe, ma mantieni i numeri come numeri
                    row_data = []
                    for cell in row:
                        if isinstance(cell.value, (int, float)):
                            row_data.append(cell.value)
                        else:
                            row_data.append(str(cell.value) if cell.value is not None else '')
                    data.append(row_data)
                    
                # Cerca formule
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and str(cell.value).startswith('='):
                            result_value = sheet_data[cell.coordinate].value
                            if result_value is None:
                                result_value = 0
                                
                            result['formulas'].append({
                                'cell': cell.coordinate,
                                'formula': str(cell.value),
                                'result': result_value
                            })
                            
                result['sheets'].append({
                    'name': sheet_name,
                    'data': data
                })
                
            return result
            
        except Exception as e:
            print(f"Error in _process_excel: {e}")
            return {'error': str(e)}

    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Processa un file di testo"""
        result = {
            'content': '',
            'lines': []
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                result['content'] = content
                result['lines'] = content.splitlines()
                
            return result
            
        except Exception as e:
            print(f"Error in _process_text: {e}")
            return {'error': str(e)}
            
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Processa un'immagine"""
        result = {
            'text': [],
            'numbers': [],
            'dates': [],
            'emails': [],
            'urls': [],
            'statistics': {},
            'metadata': {}
        }
        
        try:
            # Estrai metadati immagine
            with Image.open(file_path) as img:
                result['metadata'] = {
                    'format': img.format,
                    'mode': img.mode,
                    'dimensions': img.size,
                }
                if hasattr(img, 'info'):
                    result['metadata'].update(img.info)
            
            # Esegui OCR se necessario
            ocr_result = self.ocr.process_image(file_path)
            result.update(ocr_result)
            
            return result
            
        except Exception as e:
            print(f"Error in _process_image: {e}")
            return {'error': str(e)}

    def extract_text_from_document(self, file_path: str) -> str:
        """Estrae solo il testo da un documento"""
        try:
            result = self.process_document(file_path)
            
            if 'error' in result:
                return ''
                
            text = []
            
            # PDF
            if 'pages' in result:
                for page in result['pages']:
                    text.append(page.get('text', ''))
                    
            # Word
            elif 'paragraphs' in result:
                for para in result['paragraphs']:
                    text.append(para.get('text', ''))
                    
            # Excel
            elif 'sheets' in result:
                for sheet in result['sheets']:
                    for row in sheet['data']:
                        text.append(' '.join(row))
                        
            # Testo
            elif 'content' in result:
                text.append(result['content'])
                
            # OCR
            elif 'text' in result:
                for item in result['text']:
                    text.append(item.get('text', ''))
                    
            return '\n'.join(text)
            
        except Exception as e:
            print(f"Error in extract_text_from_document: {e}")
            return ''

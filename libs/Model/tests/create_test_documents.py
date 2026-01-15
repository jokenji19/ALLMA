import os
import cv2
import numpy as np
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF
from PIL import Image, ImageDraw, ImageFont

def create_test_directory():
    """Crea la directory per i documenti di test"""
    test_dir = "test_documents"
    os.makedirs(test_dir, exist_ok=True)
    return test_dir

def create_test_pdf(test_dir):
    """Crea un PDF di test con testo e tabelle"""
    pdf = FPDF()
    
    # Prima pagina
    pdf.add_page()
    pdf.set_font("Arial", size=16)
    pdf.cell(200, 10, txt="Test Document", ln=1, align='C')
    
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="This is a test PDF document.", ln=1, align='L')
    pdf.cell(200, 10, txt="Created on: 2025-01-21", ln=1, align='L')
    pdf.cell(200, 10, txt="Contact: test@example.com", ln=1, align='L')
    
    # Tabella
    pdf.cell(200, 10, txt="Sample Data:", ln=1, align='L')
    col_width = 40
    row_height = 10
    
    headers = ["ID", "Name", "Value"]
    data = [
        ["1", "Item A", "100"],
        ["2", "Item B", "200"],
        ["3", "Item C", "300"]
    ]
    
    # Headers
    for header in headers:
        pdf.cell(col_width, row_height, header, border=1)
    pdf.ln()
    
    # Data
    for row in data:
        for item in row:
            pdf.cell(col_width, row_height, item, border=1)
        pdf.ln()
    
    # Salva il PDF
    pdf_path = os.path.join(test_dir, "test.pdf")
    pdf.output(pdf_path)
    return pdf_path

def create_test_word(test_dir):
    """Crea un documento Word di test"""
    doc = Document()
    
    # Aggiungi titolo
    doc.add_heading('Test Document', 0)
    
    # Aggiungi paragrafi con vari tipi di dati
    doc.add_paragraph('This is a test Word document.')
    doc.add_paragraph('Date: 2025-01-21')
    doc.add_paragraph('Email: test@example.com')
    doc.add_paragraph('Phone: +1234567890')
    
    # Aggiungi una tabella
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'ID'
    header_cells[1].text = 'Name'
    header_cells[2].text = 'Value'
    
    # Data
    data = [
        ['1', 'Item A', '100'],
        ['2', 'Item B', '200'],
        ['3', 'Item C', '300']
    ]
    
    for id, name, value in data:
        row_cells = table.add_row().cells
        row_cells[0].text = id
        row_cells[1].text = name
        row_cells[2].text = value
    
    # Salva il documento
    doc_path = os.path.join(test_dir, "test.docx")
    doc.save(doc_path)
    return doc_path

def create_test_excel(test_dir):
    """Crea un foglio Excel di test"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Test Sheet"
    
    # Headers
    ws['A1'] = 'ID'
    ws['B1'] = 'Name'
    ws['C1'] = 'Value'
    ws['D1'] = 'Date'
    ws['E1'] = 'Email'
    
    # Data
    data = [
        [1, 'Item A', 100, '2025-01-21', 'test1@example.com'],
        [2, 'Item B', 200, '2025-01-22', 'test2@example.com'],
        [3, 'Item C', 300, '2025-01-23', 'test3@example.com']
    ]
    
    for row_idx, row_data in enumerate(data, start=2):
        for col_idx, value in enumerate(row_data, start=1):
            ws.cell(row=row_idx, column=col_idx, value=value)
    
    # Salva il file
    excel_path = os.path.join(test_dir, "test.xlsx")
    wb.save(excel_path)
    return excel_path

def create_test_image(test_dir):
    """Crea un'immagine di test con testo"""
    # Crea un'immagine bianca
    width = 800
    height = 400
    image = np.full((height, width, 3), 255, dtype=np.uint8)
    
    # Aggiungi testo
    font = cv2.FONT_HERSHEY_SIMPLEX
    cv2.putText(image, "Test Image", (50, 50), font, 1.5, (0, 0, 0), 2)
    cv2.putText(image, "Date: 2025-01-21", (50, 100), font, 1, (0, 0, 0), 2)
    cv2.putText(image, "Email: test@example.com", (50, 150), font, 1, (0, 0, 0), 2)
    cv2.putText(image, "Phone: +1234567890", (50, 200), font, 1, (0, 0, 0), 2)
    
    # Salva l'immagine
    image_path = os.path.join(test_dir, "test.png")
    cv2.imwrite(image_path, image)
    return image_path

def main():
    """Funzione principale"""
    # Crea la directory di test
    test_dir = create_test_directory()
    print(f"Created test directory: {test_dir}")
    
    # Crea i documenti di test
    pdf_path = create_test_pdf(test_dir)
    print(f"Created PDF: {pdf_path}")
    
    doc_path = create_test_word(test_dir)
    print(f"Created Word document: {doc_path}")
    
    excel_path = create_test_excel(test_dir)
    print(f"Created Excel file: {excel_path}")
    
    image_path = create_test_image(test_dir)
    print(f"Created image: {image_path}")
    
if __name__ == '__main__':
    main()
